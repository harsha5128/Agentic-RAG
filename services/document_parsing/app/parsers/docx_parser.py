"""
DOCX Parser for Document Parsing Service
"""

from docx import Document as DocxDocument
from io import BytesIO
from typing import Dict, Any, Tuple
import os
import tempfile

from common.observability import get_logger
from common.exceptions import DocumentParsingError, ErrorCode

from .base_parser import DocumentParser

try:
    from unstructured.partition.docx import partition_docx
except ImportError:  # pragma: no cover
    partition_docx = None

logger = get_logger(__name__)


class DOCXParser(DocumentParser):
    """Parse DOCX documents with unstructured first and python-docx fallback."""

    async def parse(self, file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from DOCX"""
        try:
            unstructured_text, unstructured_metadata = self._extract_with_unstructured(file_bytes)
            if unstructured_text:
                return unstructured_text, unstructured_metadata

            doc = DocxDocument(BytesIO(file_bytes))
            text = ""
            metadata = {
                "parser_stack": ["python-docx-fallback"],
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "headings": [],
            }

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    metadata["headings"].append(para.text)
                text += para.text + "\n"

            # Extract table contents
            for table_idx, table in enumerate(doc.tables):
                text += f"\n--- Table {table_idx + 1} ---\n"
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"

            # Extract core properties
            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "created": str(doc.core_properties.created),
                })

            return text, metadata
        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            raise DocumentParsingError(
                ErrorCode.PARSE_FAILED,
                "Failed to parse DOCX",
                {"original_error": str(e)},
                e,
            )

    def _extract_with_unstructured(self, file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        if not partition_docx:
            return "", {}

        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_file.write(file_bytes)
                temp_path = temp_file.name

            elements = partition_docx(filename=temp_path)
            text = "\n".join(str(element) for element in elements if str(element).strip())
            if not text:
                return "", {}

            return text, {
                "parser_stack": ["unstructured.partition_docx"],
                "elements": len(elements),
                "tables": sum(1 for element in elements if element.category == "Table"),
                "headings": [
                    str(element)
                    for element in elements
                    if element.category in {"Title", "Header"}
                ],
            }
        except Exception as e:
            logger.warning(f"Unstructured DOCX parsing failed, using python-docx fallback: {str(e)}")
            return "", {}
        finally:
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)
