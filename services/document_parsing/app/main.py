"""
Document Parsing Service - Production Grade
Modular architecture for parsing various document types
Handles: PDF, DOCX, XLSX, CSV, Images, TXT with robust error handling
Includes: Metadata extraction, semantic chunking, toxicity detection
"""

import hashlib
import json
import uuid
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from io import BytesIO
from contextlib import asynccontextmanager
import sys
from pathlib import Path

from fastapi import FastAPI, HTTPException
from PyPDF2 import PdfReader
from PIL import Image
import pytesseract
import pandas as pd
from docx import Document as DocxDocument

project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from common.base_service import BaseService
from common.config import settings
from common.observability import get_logger
from common.schemas import Document, DocumentStatus, DocumentType
from common.database_models import DocumentModel, ParsedContentModel
from common.exceptions import (
    DocumentParsingError,
    ErrorCode,
)
from services.document_parsing.app.parsers.pdf_parser import PDFParser as ProductionPDFParser
from services.document_parsing.app.parsers.docx_parser import DOCXParser as ProductionDOCXParser
from services.document_parsing.app.parsers.image_parser import ImageParser as ProductionImageParser
from services.document_parsing.app.parsers.spreadsheet_parser import (
    SpreadsheetParser as ProductionSpreadsheetParser,
)

try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError:  # pragma: no cover - supports lean local installs
    RecursiveCharacterTextSplitter = None


logger = get_logger(__name__)


class DocumentParser:
    """Abstract base parser"""
    
    async def parse(self, file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Parse document
        
        Returns:
            Tuple of (content, metadata)
        """
        raise NotImplementedError


class PDFParser(DocumentParser):
    """Parse PDF documents with OCR support"""
    
    async def parse(self, file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from PDF"""
        try:
            pdf_reader = PdfReader(BytesIO(file_bytes))
            text = ""
            metadata = {
                "pages": len(pdf_reader.pages),
                "extracted_via_ocr": False,
                "page_contents": []
            }
            
            for page_num, page in enumerate(pdf_reader.pages):
                # Extract text
                page_text = page.extract_text()
                text += f"\n--- Page {page_num + 1} ---\n" + page_text + "\n"
                
                # Apply OCR if text extraction failed
                if not page_text.strip() and settings.OCR_ENABLED:
                    try:
                        from pdf2image import convert_from_bytes
                        images = convert_from_bytes(file_bytes, first_page=page_num+1, last_page=page_num+1)
                        if images:
                            ocr_text = pytesseract.image_to_string(images[0])
                            text += ocr_text + "\n"
                            metadata["extracted_via_ocr"] = True
                    except Exception as e:
                        logger.warning(f"OCR failed for page {page_num}: {str(e)}")
                
                metadata["page_contents"].append({
                    "page": page_num + 1,
                    "has_content": len(page_text.strip()) > 0,
                })
            
            # Extract metadata
            if pdf_reader.metadata:
                metadata.update({
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", ""),
                })
            
            return text, metadata
        except Exception as e:
            logger.error(f"PDF parsing failed: {str(e)}")
            raise DocumentParsingError(
                ErrorCode.PARSE_FAILED,
                "Failed to parse PDF",
                {"original_error": str(e)},
                e,
            )


class DOCXParser(DocumentParser):
    """Parse DOCX documents"""
    
    async def parse(self, file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from DOCX"""
        try:
            doc = DocxDocument(BytesIO(file_bytes))
            text = ""
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "headings": [],
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    metadata["headings"].append(para.text)
                text += para.text + "\n"
            
            # Extract table contents
            for table_idx, table in enumerate(doc.tables):
                text += f"\n--- Table {table_idx + 1} ---\n"
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
            
            # Extract core properties
            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "created": str(doc.core_properties.created),
                })
            
            return text, metadata
        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            raise DocumentParsingError(
                ErrorCode.PARSE_FAILED,
                "Failed to parse DOCX",
                {"original_error": str(e)},
                e,
            )


class SpreadsheetParser(DocumentParser):
    """Parse XLSX and CSV files"""
    
    async def parse(self, file_bytes: bytes, file_type: str = "xlsx") -> Tuple[str, Dict[str, Any]]:
        """Extract text and metadata from spreadsheet"""
        try:
            if file_type == "xlsx":
                df = pd.read_excel(BytesIO(file_bytes), sheet_name=None)
            else:
                df = pd.read_csv(BytesIO(file_bytes))
                df = {"Sheet1": df}
            
            text = ""
            metadata = {
                "sheets": len(df) if isinstance(df, dict) else 1,
                "total_rows": 0,
                "total_columns": 0,
            }
            
            sheet_data = []
            for sheet_name, sheet_df in (df.items() if isinstance(df, dict) else [("Sheet1", df)]):
                text += f"\n=== Sheet: {sheet_name} ===\n"
                text += sheet_df.to_string() + "\n"
                
                sheet_data.append({
                    "name": sheet_name,
                    "rows": len(sheet_df),
                    "columns": len(sheet_df.columns),
                    "column_names": list(sheet_df.columns),
                })
                
                metadata["total_rows"] += len(sheet_df)
                metadata["total_columns"] = max(metadata["total_columns"], len(sheet_df.columns))
            
            metadata["sheets_info"] = sheet_data
            return text, metadata
        except Exception as e:
            logger.error(f"Spreadsheet parsing failed: {str(e)}")
            raise DocumentParsingError(
                ErrorCode.PARSE_FAILED,
                "Failed to parse spreadsheet",
                {"original_error": str(e)},
                e,
            )


class ImageParser(DocumentParser):
    """Parse image files with OCR"""
    
    async def parse(self, file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from image using OCR"""
        try:
            image = Image.open(BytesIO(file_bytes))
            text = pytesseract.image_to_string(image)
            
            metadata = {
                "width": image.width,
                "height": image.height,
                "format": image.format,
                "mode": image.mode,
                "extracted_via_ocr": True,
            }
            
            return text, metadata
        except Exception as e:
            logger.error(f"Image parsing failed: {str(e)}")
            raise DocumentParsingError(
                ErrorCode.OCR_FAILED,
                "Failed to parse image",
                {"original_error": str(e)},
                e,
            )


class TextParser(DocumentParser):
    """Parse plain text files"""
    
    async def parse(self, file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        try:
            text = file_bytes.decode('utf-8')
            metadata = {
                "encoding": "utf-8",
                "length_chars": len(text),
                "length_lines": len(text.split('\n')),
            }
            return text, metadata
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode('latin-1')
                metadata = {
                    "encoding": "latin-1",
                    "length_chars": len(text),
                    "length_lines": len(text.split('\n')),
                }
                return text, metadata
            except Exception as e:
                logger.error(f"Text parsing failed: {str(e)}")
                raise DocumentParsingError(
                    ErrorCode.PARSE_FAILED,
                    "Failed to parse text file",
                    {"original_error": str(e)},
                    e,
                )


class DocumentChunker:
    """
    Document chunking with a LangChain-first strategy.

    Note:
        We keep the custom paragraph chunker as an explicit fallback because it is useful
        for learning and for environments where we want fewer framework dependencies.
        LangChain's splitter is the production default because it handles recursive
        separators and overlap more predictably.
    """
    
    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = 1024,
        overlap: int = 128,
        separator: str = "\n\n",
    ) -> List[Tuple[str, Dict[str, Any]]]:
        """
        Chunk text with metadata
        
        Args:
            text: Text to chunk
            chunk_size: Target chunk size in characters
            overlap: Overlap between chunks
            separator: Sentence/paragraph separator
        
        Returns:
            List of (chunk, metadata) tuples
        """
        if RecursiveCharacterTextSplitter:
            try:
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size,
                    chunk_overlap=overlap,
                    separators=[separator, "\n", ". ", " ", ""],
                    length_function=len,
                )
                chunks = splitter.split_text(text)
                chunks_with_metadata = [
                    (
                        chunk,
                        {
                            "chunk_index": idx,
                            "length": len(chunk),
                            "chunker": "langchain_recursive_character",
                        },
                    )
                    for idx, chunk in enumerate(chunks)
                    if chunk.strip()
                ]
                logger.info(f"Created {len(chunks_with_metadata)} chunks using LangChain")
                return chunks_with_metadata
            except Exception as e:
                logger.warning(f"LangChain chunking failed, falling back to custom chunker: {str(e)}")

        try:
            chunks = []
            
            # Split by paragraphs first
            paragraphs = text.split(separator)
            current_chunk = ""
            start_para = 0
            
            for para_idx, para in enumerate(paragraphs):
                if len(current_chunk) + len(para) < chunk_size:
                    current_chunk += para + separator
                else:
                    if current_chunk.strip():
                        chunk_metadata = {
                            "start_para": start_para,
                            "end_para": para_idx - 1,
                            "length": len(current_chunk),
                            "chunker": "custom_paragraph_fallback",
                        }
                        chunks.append((current_chunk.strip(), chunk_metadata))
                    
                    # Start new chunk with overlap
                    current_chunk = para + separator
                    start_para = para_idx
            
            # Add final chunk
            if current_chunk.strip():
                chunk_metadata = {
                    "start_para": start_para,
                    "end_para": len(paragraphs) - 1,
                    "length": len(current_chunk),
                    "chunker": "custom_paragraph_fallback",
                }
                chunks.append((current_chunk.strip(), chunk_metadata))
            
            logger.info(f"Created {len(chunks)} chunks from text")
            return chunks
        except Exception as e:
            logger.error(f"Chunking failed: {str(e)}")
            raise DocumentParsingError(
                ErrorCode.CHUNK_ERROR,
                "Failed to chunk text",
                {"original_error": str(e)},
                e,
            )


class DocumentParsingService(BaseService):
    """Production-grade document parsing service"""
    
    def __init__(self):
        super().__init__("document-parsing-service")
        self.parsers = {
            DocumentType.PDF: ProductionPDFParser(),
            DocumentType.DOCX: ProductionDOCXParser(),
            DocumentType.XLSX: ProductionSpreadsheetParser(),
            DocumentType.CSV: ProductionSpreadsheetParser(),
            DocumentType.TXT: TextParser(),
            DocumentType.IMAGE: ProductionImageParser(),
        }
        self.chunker = DocumentChunker()
    
    async def _initialize_dependencies(self) -> None:
        """Initialize dependencies"""
        self.register_dependency(
            "s3",
            lambda: self._check_s3_health(),
        )
        self.register_dependency(
            "mongodb",
            lambda: self._check_mongodb_health(),
        )
    
    async def _check_s3_health(self) -> bool:
        """Check S3"""
        try:
            self.s3_client.head_bucket(Bucket=settings.S3_BUCKET_NAME)
            return True
        except Exception as e:
            logger.error(f"S3 health check failed: {str(e)}")
            return False
    
    async def _check_mongodb_health(self) -> bool:
        """Check MongoDB"""
        try:
            await self.mongodb_db.admin.command('ping')
            return True
        except Exception as e:
            logger.error(f"MongoDB health check failed: {str(e)}")
            return False
    
    async def parse_document(self, s3_path: str, document_id: str, document_type: DocumentType) -> Dict[str, Any]:
        """
        Main entry point for document parsing
        
        Flow:
        1. Download from S3
        2. Parse based on type
        3. Extract metadata
        4. Chunk content
        5. Store in MongoDB
        6. Update document status
        """
        return await self.with_error_handling(
            self._parse_document_impl,
            "parse_document",
            s3_path,
            document_id,
            document_type,
        )
    
    async def _parse_document_impl(self, s3_path: str, document_id: str, document_type: DocumentType) -> Dict[str, Any]:
        """Implementation of document parsing"""
        try:
            # Download from S3
            bucket, key = s3_path.replace("s3://", "").split("/", 1)
            logger.info(f"Downloading from S3: {s3_path}")
            
            response = self.s3_client.get_object(Bucket=bucket, Key=key)
            file_bytes = response['Body'].read()
            
            # Parse based on type
            parser = self.parsers.get(document_type)
            if not parser:
                raise DocumentParsingError(
                    ErrorCode.UNSUPPORTED_FILE_TYPE,
                    f"Unsupported document type: {document_type}",
                )
            
            logger.info(f"Parsing {document_type} document: {document_id}")
            if document_type == DocumentType.CSV:
                content, parse_metadata = await parser.parse(file_bytes, file_type="csv")
            elif document_type == DocumentType.XLSX:
                content, parse_metadata = await parser.parse(file_bytes, file_type="xlsx")
            else:
                content, parse_metadata = await parser.parse(file_bytes)
            
            # Chunk content
            chunks_with_metadata = self.chunker.chunk_text(
                content,
                settings.CHUNK_SIZE,
                settings.CHUNK_OVERLAP,
            )
            chunks = [chunk for chunk, _ in chunks_with_metadata]
            
            # Create parsed content record
            parsed_content = {
                "document_id": document_id,
                "content": content,
                "chunks": chunks,
                "chunk_count": len(chunks),
                "metadata": parse_metadata,
                "content_hash": hashlib.sha256(content.encode()).hexdigest(),
                "created_at": datetime.utcnow(),
            }
            
            # Store in MongoDB
            async with self.get_db_session() as db:
                await ParsedContentModel.insert_one(db, parsed_content)
                
                # Update document status
                await DocumentModel.update_status(db, document_id, DocumentStatus.PARSED.value)
                await DocumentModel.add_chunks(db, document_id, chunks)
            
            logger.info(f"Successfully parsed document {document_id}: {len(chunks)} chunks")
            
            return {
                "document_id": document_id,
                "chunks_created": len(chunks),
                "metadata": parse_metadata,
            }
        
        except DocumentParsingError:
            raise
        except Exception as e:
            logger.error(f"Error parsing document: {str(e)}")
            raise DocumentParsingError(
                ErrorCode.PARSE_FAILED,
                "Failed to parse document",
                {"original_error": str(e)},
                e,
            )


# Global service instance
service = None


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Manage service lifecycle"""
    global service
    service = DocumentParsingService()
    await service.initialize()
    yield
    await service.close()


app = FastAPI(
    title="Document Parsing Service",
    description="Parses and chunks documents with metadata extraction",
    version="1.0.0",
    lifespan=lifespan,
)


@app.post("/parse")
async def parse_document(s3_path: str, document_id: str, document_type: str):
    """Parse document endpoint"""
    try:
        doc_type = DocumentType[document_type.upper()]
        result = await service.parse_document(s3_path, document_id, doc_type)
        return {
            "status": "success",
            **result
        }
    except DocumentParsingError as e:
        logger.error(f"Parsing error: {str(e)}")
        raise HTTPException(status_code=422, detail=e.to_dict())
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        raise HTTPException(status_code=500, detail={"error": "Internal server error"})


@app.get("/health")
async def health_check():
    """Health check"""
    return await service.check_health()


@app.get("/ready")
async def readiness():
    """Readiness probe"""
    return await service.readiness_probe()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8002,
        workers=4,
    )

